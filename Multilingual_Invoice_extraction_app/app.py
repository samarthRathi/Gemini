from dotenv import load_dotenv
load_dotenv()

import os
import streamlit as st
from PIL import Image
import google.generativeai as genai
import pandas as pd
from io import BytesIO
from docx import Document

# ----------------------------------------
# 🎯 Gemini Vision App Class
# ----------------------------------------

class GeminiVisionApp:
    def __init__(self, task_name="image understanding", model_name="gemini-1.5-flash"):
        self.task_name = task_name
        self.model_name = model_name
        self.model = self.configure_model()
        self.default_prompt = self.get_default_prompt()

    def configure_model(self):
        api_key = os.getenv("GOOGLE_API_KEY")
        genai.configure(api_key=api_key)
        return genai.GenerativeModel(model_name=self.model_name)

    def get_default_prompt(self):
        return f"""
        You are an expert in {self.task_name}.
        Your job is to analyze the uploaded image and/or text and provide meaningful insights or structured outputs.
        If the image is blurry or incomplete, do your best. Keep your responses concise and relevant to the task.
        """

    def format_image(self, uploaded_file):
        if uploaded_file:
            return [{
                "mime_type": uploaded_file.type,
                "data": uploaded_file.getvalue()
            }]
        return [None]

    def validate_image(self, image_parts):
        check_prompt = f"Is this image relevant to the task of {self.task_name}? Just answer Yes or No."
        response = self.model.generate_content([check_prompt, image_parts[0]])
        return "yes" in response.text.strip().lower()

    def generate_response(self, user_input, image_parts, user_prompt):
        prompt_chain = [self.default_prompt]

        if user_input:
            prompt_chain.append(user_input)
        if image_parts[0]:
            prompt_chain.append(image_parts[0])
        if user_prompt:
            prompt_chain.append(user_prompt)

        response = self.model.generate_content(prompt_chain)
        return response.text

# ----------------------------------------
# 🖼️ Streamlit UI
# ----------------------------------------

st.set_page_config(page_title="Dynamic Gemini Vision App", page_icon="🧠", layout="centered")
st.title("🧠 Gemini Vision Task Assistant")

# 🧠 Task definition
st.header("1. Define Your Assistant")
user_task_name = st.text_input("Who should the model act as? (e.g., invoice parser, form analyzer, table extractor):")

# Initialize session state
if "result" not in st.session_state:
    st.session_state.result = None

# Continue if task name is given
if user_task_name:
    app = GeminiVisionApp(task_name=user_task_name)

    st.header("2. Upload Image and Enter Prompt")
    user_input = st.text_area("What would you like me to do with the image?")
    uploaded_image = st.file_uploader("Upload an image (PNG, JPG, JPEG)", type=["png", "jpg", "jpeg"])

    # Show uploaded image
    if uploaded_image and st.checkbox("Show uploaded image", value=False):
        st.image(Image.open(uploaded_image), caption="Uploaded Image", use_column_width=True)

    # Submit
    if st.button("Submit"):
        if not uploaded_image and not user_input:
            st.warning("Please provide a prompt or upload an image.")
        else:
            image_data = app.format_image(uploaded_image)

            if uploaded_image and not app.validate_image(image_data):
                st.error(f"🚫 The uploaded image does not appear relevant to '{user_task_name}'. Please upload a suitable image.")
            else:
                with st.spinner("Generating response with Gemini..."):
                    result = app.generate_response(user_input, image_data, user_prompt=None)
                    st.session_state.result = result

# Show Gemini's response
if st.session_state.result:
    st.subheader("💡 Gemini's Response")
    st.write(st.session_state.result)

    export_format = st.radio("Export extracted data as:", ["None", "Excel (.xlsx)", "Word (.docx)"], horizontal=True)

    if export_format == "Excel (.xlsx)":
        lines = [line for line in st.session_state.result.split("\n") if ":" in line]
        data = [line.split(":", 1) for line in lines]
        df = pd.DataFrame(data, columns=["Field", "Value"])

        output = BytesIO()
        df.to_excel(output, index=False)
        output.seek(0)

        st.download_button("📥 Download Excel", output, file_name="extracted_data.xlsx")

    elif export_format == "Word (.docx)":
        doc = Document()
        doc.add_heading("Extracted Data", level=1)
        for line in st.session_state.result.split("\n"):
            doc.add_paragraph(line)

        word_output = BytesIO()
        doc.save(word_output)
        word_output.seek(0)

        st.download_button("📥 Download Word", word_output, file_name="extracted_data.docx")
else:
    if user_task_name:
        st.info("📝 Submit a prompt and image to generate a response.")
    else:
        st.info("👋 Please define what kind of assistant Gemini should act as to begin.")
