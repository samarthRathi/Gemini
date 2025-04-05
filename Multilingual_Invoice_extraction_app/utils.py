import pandas as pd
from io import BytesIO
from openpyxl.styles import Font
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl import Workbook
from docx import Document
import re
import json
import os
import sqlite3
from datetime import datetime
import uuid

session_logs = {}

DB_PATH = "users.db"

def init_user_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT NOT NULL
        )
    """)
    cursor.execute("INSERT OR IGNORE INTO users (username, password) VALUES (?, ?)", ("testuser123", "123456"))
    conn.commit()
    conn.close()

init_user_db()

def verify_user(username, password):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT password FROM users WHERE username = ?", (username,))
    result = cursor.fetchone()
    conn.close()
    return result is not None and result[0] == password

# === Utility logging functions ===

def clean_text(text):
    return re.sub(r"[*_`\[\]]", "", text).strip()

def censor_text(text):
    flagged_words = ["fraud", "violence", "abuse", "drugs"]
    for word in flagged_words:
        text = re.sub(fr"\b{word}\b", "[REDACTED]", text, flags=re.IGNORECASE)
    return text

def convert_to_excel(text):
    lines = [line for line in text.split("\n") if ":" in line]
    data = [line.split(":", 1) for line in lines]
    cleaned_data = [[clean_text(k), clean_text(v)] for k, v in data]
    df = pd.DataFrame(cleaned_data, columns=["Field", "Value"])
    wb = Workbook()
    ws = wb.active
    ws.title = "Extracted Data"
    for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
        for c_idx, value in enumerate(row, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.font = Font(name='Calibri', bold=False, italic=False)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def convert_to_word(text):
    doc = Document()
    doc.add_heading("Extracted Data", level=1)
    for line in text.split("\n"):
        doc.add_paragraph(clean_text(line))
    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output

def start_session_log(user_id):
    session_id = str(uuid.uuid4())
    session_logs[user_id] = {
        "session_id": session_id,
        "session_start": datetime.now().isoformat(),
        "user_id": user_id,
        "interactions": []
    }
    return session_id

def add_session_interaction(user_id, prompt, image_name, response, unethical):
    session_logs[user_id]["interactions"].append({
        "timestamp": datetime.now().isoformat(),
        "prompt": prompt,
        "image_name": image_name,
        "response_summary": str(response)[:300],
        "flagged_unethical": unethical
    })

def end_session_log(user_id):
    if user_id not in session_logs:
        return
    session_logs[user_id]["session_end"] = datetime.now().isoformat()
    os.makedirs("logs", exist_ok=True)
    session_id = session_logs[user_id]["session_id"]
    log_path = os.path.join("logs", f"{user_id}_session_{session_id}.json")
    with open(log_path, "w") as f:
        json.dump(session_logs[user_id], f, indent=2)